import pytest
from pathlib import Path
from unittest.mock import Mock, patch, MagicMock
from documents.store import DocumentStore
from documents.ingestion import chunk_text, load_text_file, _load_pdf, _load_docx, ingest_path


@pytest.fixture
def doc_store(tmp_path):
    store = DocumentStore(persist_dir=tmp_path / "chroma")
    yield store


class TestDocumentStore:
    def test_add_and_search(self, doc_store):
        doc_store.add_documents(
            texts=["Python is a programming language", "The weather is sunny today"],
            metadatas=[{"source": "doc1.txt"}, {"source": "doc2.txt"}],
            ids=["chunk_1", "chunk_2"],
        )
        results = doc_store.search("programming", top_k=1)
        assert len(results) == 1
        assert "Python" in results[0]["text"]

    def test_search_returns_metadata(self, doc_store):
        doc_store.add_documents(
            texts=["Claude is an AI assistant"],
            metadatas=[{"source": "ai.txt", "page": 1}],
            ids=["chunk_1"],
        )
        results = doc_store.search("AI assistant", top_k=1)
        assert results[0]["metadata"]["source"] == "ai.txt"

    def test_search_empty_store(self, doc_store):
        results = doc_store.search("anything", top_k=5)
        assert results == []

    def test_search_top_k(self, doc_store):
        texts = [f"Document number {i}" for i in range(10)]
        metadatas = [{"source": f"doc_{i}.txt"} for i in range(10)]
        ids = [f"chunk_{i}" for i in range(10)]
        doc_store.add_documents(texts=texts, metadatas=metadatas, ids=ids)
        results = doc_store.search("Document", top_k=3)
        assert len(results) == 3


class TestChunking:
    def test_chunk_short_text(self):
        chunks = chunk_text("Hello world", chunk_size=100, overlap=10)
        assert len(chunks) == 1
        assert chunks[0] == "Hello world"

    def test_chunk_long_text(self):
        words = ["word"] * 200
        text = " ".join(words)
        chunks = chunk_text(text, chunk_size=50, overlap=10)
        assert len(chunks) > 1
        for chunk in chunks:
            assert len(chunk.split()) <= 50

    def test_chunk_overlap(self):
        words = [f"word{i}" for i in range(100)]
        text = " ".join(words)
        chunks = chunk_text(text, chunk_size=30, overlap=5)
        assert len(chunks) >= 2
        # Overlap means end of chunk N shares words with start of chunk N+1
        first_end_words = chunks[0].split()[-5:]
        second_start_words = chunks[1].split()[:5]
        assert first_end_words == second_start_words


class TestIngestion:
    def test_load_text_file(self, tmp_path):
        test_file = tmp_path / "test.txt"
        test_file.write_text("Hello, this is a test document.")
        text = load_text_file(test_file)
        assert text == "Hello, this is a test document."

    def test_load_markdown_file(self, tmp_path):
        test_file = tmp_path / "test.md"
        test_file.write_text("# Title\n\nSome content here.")
        text = load_text_file(test_file)
        assert "Title" in text
        assert "Some content here." in text

    def test_load_pdf_file(self, tmp_path):
        """Test PDF loading with mocked pypdf library."""
        test_file = tmp_path / "test.pdf"

        # Mock the PdfReader class
        mock_page = Mock()
        mock_page.extract_text.return_value = "Hello from PDF!\nThis is page 1."

        mock_reader = Mock()
        mock_reader.pages = [mock_page]

        with patch("pypdf.PdfReader", return_value=mock_reader):
            text = load_text_file(test_file)
            assert "Hello from PDF!" in text
            assert "This is page 1." in text

    def test_load_pdf_import_error(self, tmp_path):
        """Test that ImportError is raised when pypdf is not available."""
        test_file = tmp_path / "test.pdf"
        test_file.write_bytes(b"%PDF-1.4\n")  # Minimal PDF header

        with patch.dict("sys.modules", {"pypdf": None}):
            with pytest.raises(ImportError, match="pypdf library is required"):
                _load_pdf(test_file)

    def test_load_docx_file(self, tmp_path):
        """Test DOCX loading with real python-docx library."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        test_file = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")

        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A1"
        table.cell(0, 1).text = "B1"
        table.cell(1, 0).text = "A2"
        table.cell(1, 1).text = "B2"

        doc.save(test_file)

        text = load_text_file(test_file)
        assert "First paragraph" in text
        assert "Second paragraph" in text
        assert "A1" in text
        assert "B1" in text

    def test_load_docx_import_error(self, tmp_path):
        """Test that ImportError is raised when python-docx is not available."""
        test_file = tmp_path / "test.docx"
        test_file.write_bytes(b"PK")  # Minimal ZIP header

        with patch.dict("sys.modules", {"docx": None}):
            with pytest.raises(ImportError, match="python-docx library is required"):
                _load_docx(test_file)

    def test_load_docx_empty_handling(self, tmp_path):
        """Test DOCX with empty paragraphs."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        test_file = tmp_path / "empty.docx"
        doc = Document()
        doc.add_paragraph("")  # Empty paragraph
        doc.add_paragraph("   ")  # Whitespace only
        doc.add_paragraph("Real content")
        doc.save(test_file)

        text = load_text_file(test_file)
        assert "Real content" in text
        # Empty paragraphs should not create extra newlines
        assert text.strip() == "Real content"

    def test_load_pdf_empty_pages(self, tmp_path):
        """Test PDF with empty pages."""
        test_file = tmp_path / "empty.pdf"

        # Mock a PDF with empty pages and one page with content
        mock_empty_page = Mock()
        mock_empty_page.extract_text.return_value = "   "

        mock_content_page = Mock()
        mock_content_page.extract_text.return_value = "Real content here"

        mock_reader = Mock()
        mock_reader.pages = [mock_empty_page, mock_content_page, mock_empty_page]

        with patch("pypdf.PdfReader", return_value=mock_reader):
            text = load_text_file(test_file)
            # Empty pages should be filtered out
            assert "Real content here" in text
            assert text.count("\n") == 0  # Only one non-empty page, no newlines

    def test_ingest_pdf_end_to_end(self, tmp_path, doc_store):
        """Test ingesting a PDF file into document store."""
        test_file = tmp_path / "report.pdf"
        test_file.write_bytes(b"%PDF-1.4")  # Create the file

        # Mock PDF with multiple pages
        mock_page1 = Mock()
        mock_page1.extract_text.return_value = "Introduction to the report"

        mock_page2 = Mock()
        mock_page2.extract_text.return_value = "Chapter one contains important findings"

        mock_reader = Mock()
        mock_reader.pages = [mock_page1, mock_page2]

        with patch("pypdf.PdfReader", return_value=mock_reader):
            result = ingest_path(test_file, doc_store)
            assert "Ingested 1 file" in result

            # Search should find the ingested content
            results = doc_store.search("report", top_k=1)
            assert len(results) > 0

    def test_ingest_docx_end_to_end(self, tmp_path, doc_store):
        """Test ingesting a DOCX file into document store."""
        test_file = tmp_path / "document.docx"
        test_file.write_bytes(b"PK")  # Create the file

        # Mock DOCX
        mock_paragraph1 = Mock()
        mock_paragraph1.text = "First paragraph with data"

        mock_paragraph2 = Mock()
        mock_paragraph2.text = "Second paragraph with analysis"

        mock_doc = Mock()
        mock_doc.paragraphs = [mock_paragraph1, mock_paragraph2]
        mock_doc.tables = []

        with patch("docx.Document", return_value=mock_doc):
            result = ingest_path(test_file, doc_store)
            assert "Ingested 1 file" in result

            # Search should find the ingested content
            results = doc_store.search("analysis", top_k=1)
            assert len(results) > 0
