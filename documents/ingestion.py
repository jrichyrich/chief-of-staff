import hashlib
import logging
import os
from datetime import datetime, timezone
from pathlib import Path
from typing import TYPE_CHECKING

logger = logging.getLogger(__name__)

if TYPE_CHECKING:
    from documents.store import DocumentStore

SUPPORTED_EXTENSIONS = {".txt", ".md", ".py", ".json", ".yaml", ".yml", ".pdf", ".docx"}

EXTENSION_TO_DOC_TYPE = {
    ".md": "markdown",
    ".pdf": "pdf",
    ".docx": "docx",
    ".py": "code",
    ".json": "config",
    ".yaml": "config",
    ".yml": "config",
    ".txt": "text",
}

MAX_FILE_SIZE_BYTES = 50 * 1024 * 1024  # 50 MB


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    words = text.split()
    if len(words) <= chunk_size:
        return [text]

    chunks = []
    start = 0
    while start < len(words):
        end = start + chunk_size
        chunk = " ".join(words[start:end])
        chunks.append(chunk)
        start = end - overlap

    return chunks


def _load_pdf(file_path: Path) -> str:
    """Load text from a PDF file."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("pypdf library is required for PDF support. Install with: pip install pypdf")

    reader = PdfReader(file_path)
    text_parts = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text and page_text.strip():  # Only add non-empty pages
            text_parts.append(page_text)

    return "\n".join(text_parts)


def _load_docx(file_path: Path) -> str:
    """Load text from a DOCX file."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx library is required for DOCX support. Install with: pip install python-docx")

    doc = Document(file_path)
    text_parts = []

    # Extract paragraph text
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Extract table text
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    return "\n".join(text_parts)


def load_text_file(path: Path) -> str:
    """Load text from a file based on its extension."""
    if path.is_symlink():
        raise ValueError(f"Refusing to read symlink: {path}")

    file_size = path.stat().st_size
    if file_size > MAX_FILE_SIZE_BYTES:
        raise ValueError(f"File too large ({file_size} bytes, max {MAX_FILE_SIZE_BYTES}): {path.name}")

    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _load_pdf(path)
    elif suffix == ".docx":
        return _load_docx(path)
    else:
        return path.read_text(encoding="utf-8")


def content_hash(text: str) -> str:
    return hashlib.sha256(text.encode()).hexdigest()[:16]


def ingest_path(path: Path, document_store: "DocumentStore") -> str:
    """Ingest a file or directory into the document store.

    Returns a summary string of what was ingested.
    """
    files = []

    if path.is_file():
        files = [path]
    elif path.is_dir():
        for ext in SUPPORTED_EXTENSIONS:
            files.extend(path.glob(f"**/*{ext}"))
        # Security: reject symlinks and paths that escape the target directory
        resolved_root = path.resolve()
        safe_files = []
        for f in files:
            if f.is_symlink():
                continue
            if not f.resolve().is_relative_to(resolved_root):
                continue
            safe_files.append(f)
        files = safe_files

    if not files:
        return f"No supported files found at {path}"

    total_chunks = 0
    skipped = 0
    for file in files:
        try:
            text = load_text_file(file)
        except (ValueError, ImportError) as e:
            logger.warning("Skipping file: %s (%s)", file, e)
            skipped += 1
            continue
        except UnicodeDecodeError:
            logger.warning("Skipping file with encoding error: %s", file)
            skipped += 1
            continue
        chunks = chunk_text(text)
        file_hash = content_hash(text)

        # Lifecycle metadata for retention policies
        stat = file.stat()
        ingested_at = datetime.now(timezone.utc).isoformat()
        file_modified_at = datetime.fromtimestamp(
            os.path.getmtime(file), tz=timezone.utc
        ).isoformat()
        file_size_bytes = stat.st_size
        document_type = EXTENSION_TO_DOC_TYPE.get(file.suffix.lower(), "text")

        texts = []
        metadatas = []
        ids = []
        for i, chunk in enumerate(chunks):
            texts.append(chunk)
            metadatas.append({
                "source": str(file.name),
                "chunk_index": i,
                "created_at": ingested_at,
                "file_modified_at": file_modified_at,
                "file_size_bytes": file_size_bytes,
                "document_type": document_type,
            })
            ids.append(f"{file_hash}_{i}")

        document_store.add_documents(texts=texts, metadatas=metadatas, ids=ids)
        total_chunks += len(chunks)

    summary = f"Ingested {len(files) - skipped} file(s), {total_chunks} chunks."
    if skipped:
        summary += f" Skipped {skipped} file(s) due to size or security restrictions."
    return summary
